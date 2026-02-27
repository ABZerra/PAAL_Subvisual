#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import re
import sys
from datetime import datetime
from pathlib import Path
from typing import Any
from xml.sax.saxutils import escape
import zipfile

try:
    from docx import Document
except ImportError:  # pragma: no cover - dependency check handled at runtime
    Document = None

try:
    import yaml
except ImportError:  # pragma: no cover - optional parser
    yaml = None


PLACEHOLDER_PATTERN = re.compile(r"{{\s*([a-z0-9_]+)\s*}}")
LIST_NUMBER_PATTERN = re.compile(r"^\d+\.\s+")
VALID_FEE_TYPES = ("daily", "monthly", "fixed")

DEFAULT_OUTPUT_DIR = Path("docs-output/sow")
DEFAULT_CONTRACTOR_PROFILE = Path("project-context/sow/contractor_profile.yaml")
DEFAULT_TEMPLATE = Path("org/skills/sow-document-generator/templates/new_statement_of_work_template.md")
DEFAULT_SCHEMA = Path("org/skills/sow-document-generator/templates/field_schema.yaml")

DEFAULT_CONFIDENTIALITY_TEXT = (
    "This SOW and supporting materials contain confidential and proprietary business "
    "information of Subvisual, Lda. These materials may be printed or photocopied for "
    "use in evaluating the SOW, but are not to be shared with other parties."
)
DEFAULT_DELIVERIES_CLAUSE = (
    "A deliverable will only be considered complete when approved by the Client. "
    "The Client is entitled to three rounds of reviews and feedback on each deliverable "
    "before the End Date on 2. and the Total Cost on 7. are reviewed."
)
DEFAULT_CLIENT_RESPONSIBILITIES_TEXT = (
    "The Client should provide any required content, resources, feedback, and reviews "
    "on time as necessary."
)
DEFAULT_INVOICE_MODEL = "daily fee / fixed budget / monthly fee"
DEFAULT_INVOICE_PROCEDURES_TEXT = (
    "Invoices will be submitted at the end of every month described in the Period of "
    "Performance. Monthly invoices will reflect the number of working days executed by "
    "our team in the given month, in accordance with our internal tracking system. If "
    "needed, the Contractor will provide the Client with sufficient details to support "
    "its invoices, including timesheets for services performed and expense receipts and "
    "justifications for authorised expenses, unless otherwise agreed between the parties."
)
FIXED_PAYMENT_METHOD = "wired transfer / on-chain"
DEFAULT_FEE_TYPE = "daily"
DEFAULT_DAILY_FEE = "600€"
DEFAULT_WEEKLY_SCHEDULE = "5 days/week"
DATE_OUTPUT_FORMAT = "%d/%m/%Y"
DATE_INPUT_FORMATS = ("%d/%m/%Y", "%Y-%m-%d")


def find_repo_root(start: Path) -> Path:
    for candidate in (start, *start.parents):
        if (candidate / ".git").exists():
            return candidate
    return start


def load_structured_file(path: Path) -> dict[str, Any]:
    text = path.read_text(encoding="utf-8")

    try:
        data = json.loads(text)
    except json.JSONDecodeError:
        if yaml is None:
            raise ValueError(
                f"Unable to parse {path} as JSON and PyYAML is not installed for YAML parsing."
            ) from None
        data = yaml.safe_load(text)

    if not isinstance(data, dict):
        raise ValueError(f"Expected object/dict at top-level in {path}.")
    return data


def deep_merge(base: dict[str, Any], overrides: dict[str, Any]) -> dict[str, Any]:
    merged: dict[str, Any] = dict(base)
    for key, value in overrides.items():
        if isinstance(value, dict) and isinstance(merged.get(key), dict):
            merged[key] = deep_merge(merged[key], value)
        else:
            merged[key] = value
    return merged


def slugify(value: str) -> str:
    normalized = re.sub(r"[^a-z0-9]+", "-", value.lower()).strip("-")
    return normalized or "value"


def parse_date(value: str, field_name: str) -> datetime:
    for date_format in DATE_INPUT_FORMATS:
        try:
            return datetime.strptime(value, date_format)
        except ValueError:
            continue
    accepted = "DD/MM/YYYY or YYYY-MM-DD"
    raise ValueError(f"Invalid date for {field_name}: '{value}'. Expected {accepted}.")


def normalize_date(value: str, field_name: str) -> str:
    if value is None:
        return ""
    raw = str(value).strip()
    if not raw:
        return ""
    return parse_date(raw, field_name).strftime(DATE_OUTPUT_FORMAT)


def ensure_string_list(value: Any) -> list[str]:
    if value is None:
        return []
    if isinstance(value, list):
        return [str(item).strip() for item in value if str(item).strip()]
    if isinstance(value, str):
        return [line.strip() for line in value.splitlines() if line.strip()]
    raise ValueError("Expected list or string.")


def ensure_fee_rows(value: Any) -> list[dict[str, str]]:
    if not isinstance(value, list):
        raise ValueError("fee_schedule must be a list.")
    rows: list[dict[str, str]] = []
    for item in value:
        if not isinstance(item, dict):
            raise ValueError("Each fee_schedule item must be an object.")
        fee_type = str(item.get("fee_type", DEFAULT_FEE_TYPE)).strip().lower() or DEFAULT_FEE_TYPE
        if fee_type not in VALID_FEE_TYPES:
            choices = ", ".join(VALID_FEE_TYPES)
            raise ValueError(f"Invalid fee_type '{fee_type}'. Expected one of: {choices}.")
        fee = str(item.get("fee", "")).strip()
        if not fee and fee_type == DEFAULT_FEE_TYPE:
            fee = DEFAULT_DAILY_FEE
        row = {
            "fee_type": fee_type,
            "role": str(item.get("role", "")).strip(),
            "fee": fee,
            "allocation": str(
                item.get("allocation", item.get("schedule", ""))
            ).strip()
            or DEFAULT_WEEKLY_SCHEDULE,
            "duration": str(item.get("duration", "")).strip(),
            "estimation": str(
                item.get("estimation", item.get("cost_estimation", ""))
            ).strip(),
        }
        rows.append(row)
    return rows


def normalize_payload(raw: dict[str, Any], contractor_defaults: dict[str, Any]) -> dict[str, Any]:
    contractor_overrides = raw.get("contractor_profile_overrides", {})
    if contractor_overrides and not isinstance(contractor_overrides, dict):
        raise ValueError("contractor_profile_overrides must be an object when provided.")
    contractor = deep_merge(contractor_defaults, contractor_overrides or {})
    payment = contractor.get("payment", {})
    if not isinstance(payment, dict):
        raise ValueError("contractor profile payment block must be an object.")

    bill_to = raw.get("bill_to", {})
    if bill_to and not isinstance(bill_to, dict):
        raise ValueError("bill_to must be an object when provided.")

    execution = raw.get("execution", {})
    if execution and not isinstance(execution, dict):
        raise ValueError("execution must be an object when provided.")

    overrides = raw.get("overrides", {})
    if overrides and not isinstance(overrides, dict):
        raise ValueError("overrides must be an object when provided.")

    payload: dict[str, Any] = {
        "prepared_for": str(raw.get("prepared_for", "")).strip(),
        "date_issued": normalize_date(str(raw.get("date_issued", "")).strip(), "date_issued"),
        "client_legal_name": str(raw.get("client_legal_name", "")).strip(),
        "project_title": str(raw.get("project_title", "")).strip(),
        "project_summary": str(raw.get("project_summary", "")).strip(),
        "start_date": normalize_date(str(raw.get("start_date", "")).strip(), "start_date"),
        "end_date": normalize_date(str(raw.get("end_date", "")).strip(), "end_date"),
        "project_description": str(raw.get("project_description", "")).strip(),
        "project_cost_total": str(raw.get("project_cost_total", "")).strip(),
        "cost_resume_title": str(raw.get("cost_resume_title", "")).strip(),
        "cost_resume_value": str(raw.get("cost_resume_value", "")).strip(),
        "bill_to_client": str(
            bill_to.get("client", raw.get("bill_to_client", ""))
        ).strip(),
        "bill_to_address": str(
            bill_to.get("address", raw.get("bill_to_address", ""))
        ).strip(),
        "bill_to_email": str(
            bill_to.get("email", raw.get("bill_to_email", ""))
        ).strip(),
        "company_signing_date": normalize_date(
            execution.get("company_signing_date", raw.get("company_signing_date", "")),
            "company_signing_date",
        ),
        "client_signatory_name": str(
            execution.get("client_name", raw.get("client_signatory_name", ""))
        ).strip(),
        "client_signing_date": normalize_date(
            execution.get("client_signing_date", raw.get("client_signing_date", "")),
            "client_signing_date",
        ),
        "client_address": str(
            execution.get("client_address", raw.get("client_address", ""))
        ).strip(),
        "invoice_model": str(
            overrides.get("invoice_model", raw.get("invoice_model", DEFAULT_INVOICE_MODEL))
        ).strip()
        or DEFAULT_INVOICE_MODEL,
        "payment_method": FIXED_PAYMENT_METHOD,
        "statement_of_confidentiality": str(
            overrides.get(
                "statement_of_confidentiality",
                raw.get("statement_of_confidentiality", DEFAULT_CONFIDENTIALITY_TEXT),
            )
        ).strip()
        or DEFAULT_CONFIDENTIALITY_TEXT,
        "deliveries_clause": str(
            overrides.get("deliveries_clause", raw.get("deliveries_clause", DEFAULT_DELIVERIES_CLAUSE))
        ).strip()
        or DEFAULT_DELIVERIES_CLAUSE,
        "client_responsibilities_text": str(
            overrides.get(
                "client_responsibilities_text",
                raw.get("client_responsibilities_text", DEFAULT_CLIENT_RESPONSIBILITIES_TEXT),
            )
        ).strip()
        or DEFAULT_CLIENT_RESPONSIBILITIES_TEXT,
        "invoice_procedures_text": str(
            overrides.get(
                "invoice_procedures_text",
                raw.get("invoice_procedures_text", DEFAULT_INVOICE_PROCEDURES_TEXT),
            )
        ).strip()
        or DEFAULT_INVOICE_PROCEDURES_TEXT,
        "company_name": str(contractor.get("company_name", "")).strip(),
        "company_address": str(contractor.get("company_address", "")).strip(),
        "country": str(contractor.get("country", "")).strip(),
        "phone": str(contractor.get("phone", "")).strip(),
        "email": str(contractor.get("email", "")).strip(),
        "website": str(contractor.get("website", "")).strip(),
        "prepared_by_name": str(
            raw.get("prepared_by_name", contractor.get("prepared_by_name", ""))
        ).strip(),
        "prepared_by_role": str(
            raw.get("prepared_by_role", contractor.get("prepared_by_role", ""))
        ).strip(),
        "payment_company_name": str(
            payment.get("company_name", contractor.get("company_name", ""))
        ).strip(),
        "payment_bank_name": str(payment.get("bank_name", "")).strip(),
        "payment_iban": str(payment.get("iban", "")).strip(),
        "payment_swift": str(payment.get("swift", "")).strip(),
        "payment_wallet": str(payment.get("wallet", "")).strip(),
    }

    payload["non_working_days"] = ensure_string_list(raw.get("non_working_days", []))
    payload["deliverables"] = ensure_string_list(raw.get("deliverables", []))
    payload["fee_schedule"] = ensure_fee_rows(raw.get("fee_schedule", []))

    if not payload["cost_resume_title"] and payload["project_title"]:
        payload["cost_resume_title"] = payload["project_title"]
    if not payload["cost_resume_value"] and payload["project_cost_total"]:
        payload["cost_resume_value"] = payload["project_cost_total"]

    return payload


def validate_payload(payload: dict[str, Any], schema: dict[str, Any]) -> None:
    errors: list[str] = []
    required_fields = schema.get("required_fields", [])
    if not isinstance(required_fields, list):
        raise ValueError("field_schema required_fields must be a list.")

    for field_name in required_fields:
        value = payload.get(field_name)
        if value is None:
            errors.append(f"Missing required field: {field_name}")
            continue
        if isinstance(value, str) and not value.strip():
            errors.append(f"Missing required field: {field_name}")
        if isinstance(value, list) and len(value) == 0:
            errors.append(f"Missing required field: {field_name}")

    if payload.get("start_date") and payload.get("end_date"):
        start_date = parse_date(payload["start_date"], "start_date")
        end_date = parse_date(payload["end_date"], "end_date")
        if start_date > end_date:
            errors.append("Invalid date range: start_date must be <= end_date.")

    fee_schedule = payload.get("fee_schedule", [])
    if not isinstance(fee_schedule, list) or len(fee_schedule) == 0:
        errors.append("fee_schedule must contain at least one row.")
    else:
        required_fee_row_fields = schema.get("required_fee_row_fields", [])
        if not isinstance(required_fee_row_fields, list):
            raise ValueError("field_schema required_fee_row_fields must be a list.")
        for index, row in enumerate(fee_schedule, start=1):
            if not isinstance(row, dict):
                errors.append(f"fee_schedule row {index} must be an object.")
                continue
            for field_name in required_fee_row_fields:
                if not str(row.get(field_name, "")).strip():
                    errors.append(f"fee_schedule row {index} missing required field: {field_name}")

    if errors:
        raise ValueError("\n".join(errors))


def format_bullet_list(items: list[str]) -> str:
    if not items:
        return "-"
    return "\n".join(f"- {item}" for item in items)


def format_numbered_list(items: list[str]) -> str:
    return "\n".join(f"{index}. {item}" for index, item in enumerate(items, start=1))


def parse_number_with_units(raw_value: str) -> tuple[float, str, str, int] | None:
    text = str(raw_value).strip()
    if not text:
        return None
    match = re.search(r"-?\d+(?:[.,]\d+)?", text)
    if not match:
        return None
    number_token = match.group(0)
    decimals = 0
    if "." in number_token:
        decimals = len(number_token.split(".", 1)[1])
    elif "," in number_token:
        decimals = len(number_token.split(",", 1)[1])
    number = float(number_token.replace(",", "."))
    prefix = text[: match.start()].strip()
    suffix = text[match.end() :].strip()
    return number, prefix, suffix, decimals


def format_numeric_with_units(total: float, prefix: str, suffix: str, decimals: int) -> str:
    if decimals <= 0:
        number_text = str(int(round(total)))
    else:
        number_text = f"{total:.{decimals}f}".rstrip("0").rstrip(".")

    if prefix and suffix:
        return f"{prefix}{number_text} {suffix}"
    if prefix:
        return f"{prefix}{number_text}"
    if suffix:
        if suffix in {"€", "$", "£", "%"}:
            return f"{number_text}{suffix}"
        return f"{number_text} {suffix}"
    return number_text


def sum_column(values: list[str]) -> str:
    parsed = [parse_number_with_units(value) for value in values]
    if not parsed or any(item is None for item in parsed):
        return "-"

    first = parsed[0]
    assert first is not None
    _, prefix, suffix, decimals = first
    total = 0.0
    max_decimals = decimals

    for item in parsed:
        assert item is not None
        number, current_prefix, current_suffix, current_decimals = item
        if current_prefix != prefix or current_suffix != suffix:
            return "-"
        total += number
        max_decimals = max(max_decimals, current_decimals)

    return format_numeric_with_units(total, prefix, suffix, max_decimals)


def format_fee_schedule_rows(rows: list[dict[str, str]]) -> str:
    rendered = []
    for row in rows:
        rendered.append(
            "| {role} | {fee} | {allocation} | {duration} | {estimation} |".format(
                role=row.get("role", ""),
                fee=row.get("fee", ""),
                allocation=row.get("allocation", ""),
                duration=row.get("duration", ""),
                estimation=row.get("estimation", ""),
            )
        )
    return "\n".join(rendered)


def format_fee_totals_row(rows: list[dict[str, str]]) -> str:
    if not rows:
        return "| Totals | - | - | - | - |"

    fee_total = sum_column([row.get("fee", "") for row in rows])
    allocation_total = sum_column([row.get("allocation", "") for row in rows])
    duration_total = sum_column([row.get("duration", "") for row in rows])
    estimation_total = sum_column([row.get("estimation", "") for row in rows])
    return f"| Totals | {fee_total} | {allocation_total} | {duration_total} | {estimation_total} |"


def build_template_values(payload: dict[str, Any]) -> dict[str, str]:
    values = {key: str(value) for key, value in payload.items() if isinstance(value, str)}
    values["non_working_days_bullets"] = format_bullet_list(payload.get("non_working_days", []))
    values["deliverables_numbered"] = format_numbered_list(payload.get("deliverables", []))
    fee_rows = payload.get("fee_schedule", [])
    values["fee_schedule_rows"] = format_fee_schedule_rows(fee_rows)
    values["fee_totals_row"] = format_fee_totals_row(fee_rows)
    return values


def render_markdown(template_text: str, payload: dict[str, Any]) -> str:
    values = build_template_values(payload)

    def replace(match: re.Match[str]) -> str:
        key = match.group(1)
        return values.get(key, match.group(0))

    return PLACEHOLDER_PATTERN.sub(replace, template_text)


def ensure_no_unresolved_placeholders(markdown_text: str) -> None:
    unresolved = sorted(set(PLACEHOLDER_PATTERN.findall(markdown_text)))
    if unresolved:
        names = ", ".join(unresolved)
        raise ValueError(f"Unresolved template placeholders: {names}")


def is_table_separator(line: str) -> bool:
    stripped = line.strip()
    if not stripped.startswith("|"):
        return False
    cells = [cell.strip() for cell in stripped.strip("|").split("|")]
    if not cells:
        return False
    return all(bool(re.fullmatch(r":?-{3,}:?", cell.replace(" ", ""))) for cell in cells)


def parse_table_cells(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def add_table_to_docx(document: Any, header: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1 + len(rows), cols=len(header))
    for col, value in enumerate(header):
        table.cell(0, col).text = value
    for row_index, row in enumerate(rows, start=1):
        for col in range(len(header)):
            value = row[col] if col < len(row) else ""
            table.cell(row_index, col).text = value
    document.add_paragraph("")


def create_docx_from_markdown(markdown_text: str, output_path: Path) -> None:
    if Document is None:
        create_fallback_docx(markdown_text, output_path)
        return

    document = Document()
    lines = markdown_text.splitlines()
    index = 0

    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()

        if not stripped:
            document.add_paragraph("")
            index += 1
            continue

        if stripped.startswith("|") and (index + 1) < len(lines) and is_table_separator(lines[index + 1]):
            header = parse_table_cells(lines[index])
            index += 2
            table_rows: list[list[str]] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_rows.append(parse_table_cells(lines[index]))
                index += 1
            add_table_to_docx(document, header, table_rows)
            continue

        if stripped.startswith("### "):
            document.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith("## "):
            document.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("# "):
            document.add_heading(stripped[2:].strip(), level=1)
        elif stripped == "---":
            document.add_paragraph("-" * 32)
        elif LIST_NUMBER_PATTERN.match(stripped):
            document.add_paragraph(LIST_NUMBER_PATTERN.sub("", stripped), style="List Number")
        elif stripped.startswith("- "):
            document.add_paragraph(stripped[2:].strip(), style="List Bullet")
        else:
            document.add_paragraph(line)
        index += 1

    document.save(str(output_path))


def create_fallback_docx(markdown_text: str, output_path: Path) -> None:
    paragraphs: list[str] = []
    for raw_line in markdown_text.splitlines():
        text = escape(raw_line)
        if text:
            paragraphs.append(
                f'<w:p><w:r><w:t xml:space="preserve">{text}</w:t></w:r></w:p>'
            )
        else:
            paragraphs.append("<w:p/>")

    content_types_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>
"""
    rels_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>
"""
    document_rels_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"></Relationships>
"""
    document_xml = (
        """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>"""
        """<w:document xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas" """
        """xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006" """
        """xmlns:o="urn:schemas-microsoft-com:office:office" """
        """xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" """
        """xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" """
        """xmlns:v="urn:schemas-microsoft-com:vml" """
        """xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing" """
        """xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" """
        """xmlns:w10="urn:schemas-microsoft-com:office:word" """
        """xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" """
        """xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml" """
        """xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup" """
        """xmlns:wpi="http://schemas.microsoft.com/office/word/2010/wordprocessingInk" """
        """xmlns:wne="http://schemas.microsoft.com/office/2006/wordml" """
        """xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape" """
        """mc:Ignorable="w14 wp14"><w:body>"""
        + "".join(paragraphs)
        + """<w:sectPr><w:pgSz w:w="12240" w:h="15840"/><w:pgMar w:top="1440" w:right="1440" """
        """w:bottom="1440" w:left="1440" w:header="708" w:footer="708" w:gutter="0"/></w:sectPr>"""
        """</w:body></w:document>"""
    )

    with zipfile.ZipFile(output_path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", content_types_xml)
        archive.writestr("_rels/.rels", rels_xml)
        archive.writestr("word/document.xml", document_xml)
        archive.writestr("word/_rels/document.xml.rels", document_rels_xml)


def render_and_validate(template_text: str, payload: dict[str, Any], schema: dict[str, Any]) -> str:
    validate_payload(payload, schema)
    markdown = render_markdown(template_text, payload)
    ensure_no_unresolved_placeholders(markdown)
    return markdown


def output_basename(payload: dict[str, Any]) -> str:
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    client_seed = payload.get("client_legal_name") or payload.get("bill_to_client") or "client"
    project_seed = payload.get("project_title") or "project"
    return f"{timestamp}_{slugify(client_seed)}_{slugify(project_seed)}_sow"


def write_outputs(markdown: str, payload: dict[str, Any], output_dir: Path) -> tuple[Path, Path]:
    output_dir.mkdir(parents=True, exist_ok=True)
    base = output_basename(payload)
    markdown_path = output_dir / f"{base}.md"
    docx_path = output_dir / f"{base}.docx"

    markdown_path.write_text(markdown, encoding="utf-8")
    create_docx_from_markdown(markdown, docx_path)
    return markdown_path, docx_path


def prompt_text(prompt: str, required: bool = False, default: str | None = None) -> str:
    while True:
        suffix = f" [{default}]" if default else ""
        value = input(f"{prompt}{suffix}: ").strip()
        if value:
            return value
        if default is not None:
            return default
        if not required:
            return ""
        print("Value is required.")


def prompt_date(prompt: str, default: str | None = None) -> str:
    while True:
        value = prompt_text(prompt, required=True, default=default)
        try:
            return normalize_date(value, prompt)
        except ValueError as exc:
            print(str(exc))


def prompt_yes_no(prompt: str, default: bool = False) -> bool:
    suffix = " [Y/n]" if default else " [y/N]"
    while True:
        answer = input(f"{prompt}{suffix}: ").strip().lower()
        if not answer:
            return default
        if answer in {"y", "yes"}:
            return True
        if answer in {"n", "no"}:
            return False
        print("Please enter y or n.")


def prompt_multiline(prompt: str, default: str) -> str:
    print(f"{prompt} (finish with empty line; leave first line empty to keep default)")
    lines: list[str] = []
    while True:
        line = input()
        if not line:
            break
        lines.append(line)
    if not lines:
        return default
    return "\n".join(lines)


def prompt_choice(prompt: str, choices: tuple[str, ...], default: str) -> str:
    options = "/".join(choices)
    while True:
        value = prompt_text(f"{prompt} ({options})", default=default).lower()
        if value in choices:
            return value
        print(f"Please choose one of: {options}.")


def maybe_prompt_text(
    prompt: str,
    value: Any,
    required: bool,
    only_missing: bool,
    fallback_default: str | None = None,
) -> str:
    existing = "" if value is None else str(value).strip()
    if only_missing and existing:
        return existing
    default = existing or fallback_default
    return prompt_text(prompt, required=required, default=default if default else None)


def maybe_prompt_date(prompt: str, value: Any, only_missing: bool) -> str:
    existing = "" if value is None else str(value).strip()
    if only_missing and existing:
        return existing
    return prompt_date(prompt, default=existing or None)


def prompt_list(
    prompt: str,
    required: bool = False,
    default_items: list[str] | None = None,
    only_missing: bool = False,
) -> list[str]:
    if default_items:
        if only_missing or prompt_yes_no(f"Use existing values for {prompt}?", default=True):
            return default_items
    print(f"{prompt} (one per line, finish with empty line)")
    items: list[str] = []
    while True:
        value = input("- ").strip()
        if not value:
            break
        items.append(value)
    if required and not items:
        print("At least one item is required.")
        return prompt_list(prompt, required=required, default_items=default_items, only_missing=only_missing)
    return items


def prompt_fee_schedule_rows(
    default_rows: list[dict[str, Any]] | None = None,
    only_missing: bool = False,
) -> list[dict[str, str]]:
    if default_rows:
        try:
            normalized_defaults = ensure_fee_rows(default_rows)
        except ValueError:
            normalized_defaults = []
        if normalized_defaults and (only_missing or prompt_yes_no("Use existing fee schedule rows?", default=True)):
            return normalized_defaults

    while True:
        raw_count = input("How many fee schedule rows? [min 1]: ").strip()
        if not raw_count.isdigit() or int(raw_count) < 1:
            print("Please enter a number >= 1.")
            continue
        count = int(raw_count)
        break

    rows: list[dict[str, str]] = []
    for index in range(1, count + 1):
        print(f"Fee row {index}")
        fee_type = prompt_choice("  Fee type", VALID_FEE_TYPES, DEFAULT_FEE_TYPE)
        rows.append(
            {
                "fee_type": fee_type,
                "role": prompt_text("  Role", required=True),
                "fee": prompt_text(
                    "  Fee",
                    required=fee_type != DEFAULT_FEE_TYPE,
                    default=DEFAULT_DAILY_FEE if fee_type == DEFAULT_FEE_TYPE else None,
                ),
                "allocation": prompt_text(
                    "  Allocation",
                    required=True,
                    default=DEFAULT_WEEKLY_SCHEDULE,
                ),
                "duration": prompt_text("  Duration", required=True),
                "estimation": prompt_text("  Estimation", required=True),
            }
        )
    print("\nFee Schedule Table Preview")
    print("| Role | Fee | Allocation | duration | Estimation |")
    print("| --- | --- | --- | --- | --- |")
    print(format_fee_schedule_rows(rows))
    print(format_fee_totals_row(rows))
    return rows


def collect_contractor_overrides(defaults: dict[str, Any]) -> dict[str, Any]:
    if not prompt_yes_no("Override contractor defaults for this run?", default=False):
        return {}

    payment_defaults = defaults.get("payment", {})
    overrides = {
        "company_name": prompt_text("Company name", default=str(defaults.get("company_name", ""))),
        "company_address": prompt_text(
            "Company address", default=str(defaults.get("company_address", ""))
        ),
        "country": prompt_text("Country", default=str(defaults.get("country", ""))),
        "phone": prompt_text("Phone", default=str(defaults.get("phone", ""))),
        "email": prompt_text("Email", default=str(defaults.get("email", ""))),
        "website": prompt_text("Website", default=str(defaults.get("website", ""))),
        "prepared_by_name": prompt_text(
            "Prepared by (name)", default=str(defaults.get("prepared_by_name", ""))
        ),
        "prepared_by_role": prompt_text(
            "Prepared by (role)", default=str(defaults.get("prepared_by_role", ""))
        ),
        "payment": {
            "company_name": prompt_text(
                "Payment company name", default=str(payment_defaults.get("company_name", ""))
            ),
            "bank_name": prompt_text(
                "Payment bank name", default=str(payment_defaults.get("bank_name", ""))
            ),
            "iban": prompt_text("Payment IBAN", default=str(payment_defaults.get("iban", ""))),
            "swift": prompt_text("Payment BIC/SWIFT", default=str(payment_defaults.get("swift", ""))),
            "wallet": prompt_text("Payment wallet", default=str(payment_defaults.get("wallet", ""))),
        },
    }
    return overrides


def collect_interactive_input(
    contractor_defaults: dict[str, Any],
    schema: dict[str, Any],
    prefill: dict[str, Any] | None = None,
    only_missing: bool = False,
) -> dict[str, Any]:
    source = prefill or {}
    source_overrides = source.get("overrides", {}) if isinstance(source.get("overrides", {}), dict) else {}
    source_bill_to = source.get("bill_to", {}) if isinstance(source.get("bill_to", {}), dict) else {}
    source_execution = source.get("execution", {}) if isinstance(source.get("execution", {}), dict) else {}
    raw: dict[str, Any] = {
        "overrides": dict(source_overrides),
        "bill_to": dict(source_bill_to),
        "execution": dict(source_execution),
    }
    section_names = [s.get("name") for s in schema.get("sections", []) if isinstance(s, dict)]

    if "contractor_profile" in section_names:
        print("\n== Contractor Profile ==")
        if only_missing:
            raw["contractor_profile_overrides"] = source.get("contractor_profile_overrides", {})
        else:
            raw["contractor_profile_overrides"] = collect_contractor_overrides(contractor_defaults)

    if "prepared_for_and_issue" in section_names:
        print("\n== Prepared For and Date Issued ==")
        raw["prepared_for"] = maybe_prompt_text(
            "Prepared For",
            source.get("prepared_for", ""),
            required=True,
            only_missing=only_missing,
        )
        raw["date_issued"] = maybe_prompt_date(
            "Date Issued (DD/MM/YYYY)",
            source.get("date_issued", ""),
            only_missing=only_missing,
        )

        prepared_by_default = str(source.get("prepared_by_name", "")).strip() or str(
            contractor_defaults.get("prepared_by_name", "")
        )
        prepared_role_default = str(source.get("prepared_by_role", "")).strip() or str(
            contractor_defaults.get("prepared_by_role", "")
        )
        raw["prepared_by_name"] = maybe_prompt_text(
            "Prepared By name",
            source.get("prepared_by_name", ""),
            required=True,
            only_missing=only_missing,
            fallback_default=prepared_by_default,
        )
        raw["prepared_by_role"] = maybe_prompt_text(
            "Prepared By role",
            source.get("prepared_by_role", ""),
            required=True,
            only_missing=only_missing,
            fallback_default=prepared_role_default,
        )

    if "introduction_background" in section_names:
        print("\n== Introduction and Background ==")
        raw["client_legal_name"] = maybe_prompt_text(
            "Client legal name",
            source.get("client_legal_name", ""),
            required=True,
            only_missing=only_missing,
        )
        raw["project_title"] = maybe_prompt_text(
            "Project title",
            source.get("project_title", ""),
            required=True,
            only_missing=only_missing,
        )
        raw["project_summary"] = maybe_prompt_text(
            "Project summary sentence (used in section 1)",
            source.get("project_summary", ""),
            required=True,
            only_missing=only_missing,
        )

    if "period_of_performance" in section_names:
        print("\n== Period of Performance ==")
        raw["start_date"] = maybe_prompt_date(
            "Start Date (DD/MM/YYYY)",
            source.get("start_date", ""),
            only_missing=only_missing,
        )
        raw["end_date"] = maybe_prompt_date(
            "End Date (DD/MM/YYYY)",
            source.get("end_date", ""),
            only_missing=only_missing,
        )
        default_non_working_days = ensure_string_list(source.get("non_working_days", []))
        raw["non_working_days"] = prompt_list(
            "Non-working days (optional, use DD/MM/YYYY where possible)",
            required=False,
            default_items=default_non_working_days,
            only_missing=only_missing,
        )

    if "scope_of_work" in section_names:
        print("\n== Scope of Work ==")
        raw["project_description"] = maybe_prompt_text(
            "Project description",
            source.get("project_description", ""),
            required=True,
            only_missing=only_missing,
        )
        default_deliverables = ensure_string_list(source.get("deliverables", []))
        raw["deliverables"] = prompt_list(
            "Deliverables",
            required=True,
            default_items=default_deliverables,
            only_missing=only_missing,
        )
        if not only_missing and prompt_yes_no("Customize Deliveries clause?", default=False):
            raw["overrides"]["deliveries_clause"] = prompt_multiline(
                "Enter Deliveries clause",
                DEFAULT_DELIVERIES_CLAUSE,
            )

    if "fee_schedule" in section_names:
        print("\n== Fee Schedule ==")
        raw["fee_schedule"] = prompt_fee_schedule_rows(
            default_rows=source.get("fee_schedule") if isinstance(source.get("fee_schedule"), list) else None,
            only_missing=only_missing,
        )
        raw["project_cost_total"] = maybe_prompt_text(
            "Project costs total",
            source.get("project_cost_total", ""),
            required=True,
            only_missing=only_missing,
        )
        raw["cost_resume_title"] = maybe_prompt_text(
            "Costs resume title",
            source.get("cost_resume_title", ""),
            required=True,
            only_missing=only_missing,
            fallback_default=str(source.get("project_title", "")).strip() or None,
        )
        raw["cost_resume_value"] = maybe_prompt_text(
            "Costs resume value",
            source.get("cost_resume_value", ""),
            required=True,
            only_missing=only_missing,
            fallback_default=str(source.get("project_cost_total", "")).strip() or None,
        )

    if "bill_to" in section_names:
        print("\n== Bill To ==")
        raw["bill_to"] = {
            "client": maybe_prompt_text(
                "Bill To client",
                source_bill_to.get("client", source.get("bill_to_client", "")),
                required=True,
                only_missing=only_missing,
            ),
            "address": maybe_prompt_text(
                "Bill To address",
                source_bill_to.get("address", source.get("bill_to_address", "")),
                required=True,
                only_missing=only_missing,
            ),
            "email": maybe_prompt_text(
                "Bill To email",
                source_bill_to.get("email", source.get("bill_to_email", "")),
                required=True,
                only_missing=only_missing,
            ),
        }

    if "legal_boilerplate" in section_names and not only_missing:
        print("\n== Legal and Invoice Boilerplate ==")
        if prompt_yes_no("Customize statement of confidentiality?", default=False):
            raw["overrides"]["statement_of_confidentiality"] = prompt_multiline(
                "Enter Statement of Confidentiality text",
                DEFAULT_CONFIDENTIALITY_TEXT,
            )
        if prompt_yes_no("Customize client responsibilities?", default=False):
            raw["overrides"]["client_responsibilities_text"] = prompt_multiline(
                "Enter Client Responsibilities text",
                DEFAULT_CLIENT_RESPONSIBILITIES_TEXT,
            )
        if prompt_yes_no("Customize invoice model line?", default=False):
            raw["overrides"]["invoice_model"] = prompt_text(
                "Invoice model line",
                required=True,
                default=DEFAULT_INVOICE_MODEL,
            )
        if prompt_yes_no("Customize invoice procedures?", default=False):
            raw["overrides"]["invoice_procedures_text"] = prompt_multiline(
                "Enter Invoice Procedures text",
                DEFAULT_INVOICE_PROCEDURES_TEXT,
            )

    if "execution" in section_names:
        print("\n== Execution and Signatures ==")
        raw["execution"] = {
            "company_signing_date": maybe_prompt_date(
                "Company signing date (DD/MM/YYYY)",
                source_execution.get("company_signing_date", source.get("company_signing_date", "")),
                only_missing=only_missing,
            ),
            "client_name": maybe_prompt_text(
                "Client signatory name",
                source_execution.get("client_name", source.get("client_signatory_name", "")),
                required=True,
                only_missing=only_missing,
            ),
            "client_signing_date": maybe_prompt_date(
                "Client signing date (DD/MM/YYYY)",
                source_execution.get("client_signing_date", source.get("client_signing_date", "")),
                only_missing=only_missing,
            ),
            "client_address": maybe_prompt_text(
                "Client signatory address",
                source_execution.get("client_address", source.get("client_address", "")),
                required=True,
                only_missing=only_missing,
            ),
        }

    return raw


def build_arg_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Generate Statement of Work documents from the project template."
    )
    parser.add_argument("--input", type=Path, help="Input YAML/JSON file with SOW data.")
    parser.add_argument(
        "--interactive",
        action="store_true",
        help="Collect required SOW fields section-by-section.",
    )
    parser.add_argument(
        "--output-dir",
        type=Path,
        default=None,
        help="Output directory for generated .md and .docx files (default: docs-output/sow).",
    )
    parser.add_argument(
        "--contractor-profile",
        type=Path,
        default=None,
        help="Path to contractor profile YAML/JSON.",
    )
    parser.add_argument(
        "--template",
        type=Path,
        default=None,
        help="Path to markdown template.",
    )
    parser.add_argument(
        "--field-schema",
        type=Path,
        default=None,
        help="Path to field schema YAML/JSON.",
    )
    return parser


def main() -> int:
    parser = build_arg_parser()
    args = parser.parse_args()

    if bool(args.input) == bool(args.interactive):
        parser.error("Use exactly one mode: either --interactive or --input <file>.")

    script_path = Path(__file__).resolve()
    repo_root = find_repo_root(script_path)

    contractor_profile_path = (
        args.contractor_profile if args.contractor_profile else repo_root / DEFAULT_CONTRACTOR_PROFILE
    )
    template_path = args.template if args.template else repo_root / DEFAULT_TEMPLATE
    schema_path = args.field_schema if args.field_schema else repo_root / DEFAULT_SCHEMA
    output_dir = args.output_dir if args.output_dir else repo_root / DEFAULT_OUTPUT_DIR

    try:
        contractor_profile = load_structured_file(contractor_profile_path)
        schema = load_structured_file(schema_path)
        template_text = template_path.read_text(encoding="utf-8")

        if args.interactive:
            raw_payload = collect_interactive_input(contractor_profile, schema)
        else:
            raw_payload = load_structured_file(args.input)
            # Input mode still offers a quick completion flow for missing details.
            if sys.stdin.isatty():
                raw_payload = collect_interactive_input(
                    contractor_profile,
                    schema,
                    prefill=raw_payload,
                    only_missing=True,
                )

        payload = normalize_payload(raw_payload, contractor_profile)
        markdown = render_and_validate(template_text, payload, schema)
        markdown_path, docx_path = write_outputs(markdown, payload, output_dir)
    except Exception as exc:  # noqa: BLE001
        print(f"Error: {exc}", file=sys.stderr)
        return 1

    print(f"Generated markdown: {markdown_path}")
    print(f"Generated docx: {docx_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
